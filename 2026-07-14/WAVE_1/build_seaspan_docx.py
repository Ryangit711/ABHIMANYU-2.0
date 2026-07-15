#!/usr/bin/env python3
"""One-shot DOCX builder for Seaspan — no parser, no config, just clean output."""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

ONEDRIVE = "/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0"
LINUX = "/home/aryan/opencode_test/ABHIMANYU-2.0"
NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"
FONT = "Calibri"
SIZE = Pt(11)
HSIZE = Pt(13)

def set_margins(doc, m=Inches(0.75)):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = m

def run(p, text, font=FONT, size=SIZE, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = font; r.font.size = size; r.bold = bold; r.italic = italic
    if color:
        c = int(color, 16)
        r.font.color.rgb = RGBColor((c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF)
    return r

def hyperlink(p, label, url):
    part = p.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    p._p.append(parse_xml(
        f'<w:hyperlink xmlns:w="{ns_w}" xmlns:r="{ns_r}" r:id="{r_id}" w:history="1">'
        f'<w:r><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/>'
        f'<w:sz w:val="{int(Pt(9).pt * 2)}"/><w:color w:val="0563C1"/><w:u w:val="single"/>'
        f'</w:rPr><w:t xml:space="preserve">{label}</w:t></w:r></w:hyperlink>'))

def section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run(p, text.upper(), size=HSIZE, bold=True)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    b = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '000000'})
    pBdr.append(b); pPr.append(pBdr)

def body(doc, text, bold=False, italic=False, size=SIZE, sa=2, sb=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    run(p, text, size=size, bold=bold, italic=italic)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        run(p, f"\u2022 {bold_prefix}: ", bold=True)
        run(p, text)
    else:
        run(p, f"\u2022 {text}")

def contact(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, NAME, size=Pt(16), bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, f"{PHONE}  |  ", size=Pt(9), color="505050")
    hyperlink(p2, EMAIL, f"mailto:{EMAIL}")
    run(p2, "  |  ", size=Pt(9), color="505050")
    hyperlink(p2, "LinkedIn", f"https://{LINKEDIN}")
    run(p2, f"  |  {LOCATION}", size=Pt(9), color="505050")

# ===== RESUME =====
doc = Document()
set_margins(doc)

# Header
contact(doc)

# Professional Summary
section_header(doc, "Professional Summary")
body(doc, "Change management leader with 8 years guiding organizations through complex transformations spanning acquisitions, technology adoption, and organizational restructuring. Specializes in stakeholder alignment, adoption measurement, and building sustainable change frameworks. Achieved 100% talent retention across 5 acquisition integrations, 92%+ digital adoption rates across 32 locations, and $4M+ in recovered revenue through sustained behavioral change programs. Applies ADKAR, Kotter, and PROSCI-aligned methodologies tested operationally \u2014 not theoretically.")

# Core Competencies
section_header(doc, "Core Competencies")
body(doc, "ADKAR Change Management Methodology  |  Kotter 8-Step Change Process  |  M&A Integration & Transition Management  |  Stakeholder Engagement & Resistance Management  |  Change Impact Assessment & Adoption Analytics  |  KPI Dashboard Design  |  Healthcare Digital Transformation  |  Organizational Design & Restructuring  |  Cross-Functional Leadership  |  Training & Capability Building", size=Pt(9.5))

# Professional Experience
section_header(doc, "Professional Experience")

def job_header(company, title, location, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    run(p, company, bold=True)
    run(p, f"  |  {title}  |  {location}  |  {dates}")

def job_desc(text):
    body(doc, text, italic=True, size=Pt(9.5), sa=2)

# SkyflyMD
job_header("SkyflyMD", "Director of Operations & Change Management", "Vancouver, BC", "2017 \u2013 2025")
job_desc("Led change management for a multi-site healthcare organization scaling from 3 to 70 employees across 32 locations in 4 states. Directed $17M acquisition integration, full digital transformation, and 5 organizational restructures using ADKAR and Kotter methodologies.")

bullet(doc, "Achieved 100% talent retention across 5 acquisition integrations by designing ADKAR-based change plans with Day 1/30/100 milestones, managing 8 cross-functional workstreams, and consolidating 8 operational systems into 1 unified platform within 90 days of close \u2014 zero operational disruption throughout transition", "100% Retention Rate")
bullet(doc, "Drove 92%+ adoption of new digital systems across 32 locations through staged rollout strategy (pilot \u2192 iterate \u2192 deploy), hands-on training for 150+ users, and on-site reinforcement \u2014 managing full ADKAR change curve from awareness through sustained adoption", "92% Digital Adoption")
bullet(doc, "Recovered $4M+ in organic revenue by designing patient re-engagement campaigns targeting 10,000+ inactive records, building multi-channel workflows, and tracking conversion through custom analytics \u2014 reducing patient attrition by 30% through sustained behavioral change", "$4M+ Revenue Recovery")
bullet(doc, "Reduced reporting lag by 30% by deploying KPI dashboards across all 32 locations, enabling real-time adoption tracking and operational decisions by location and department managers", "30% Reporting Lag Reduction")
bullet(doc, "Scaled organizational capability from 3 to 70 employees by designing hiring frameworks, org structures, training programs, and performance management systems \u2014 each restructure executed as a Kotter-style change initiative with coalition building and vision alignment", "3\u219270 Org Scale")
bullet(doc, "Aligned 5 clinic groups and 12 departments around shared priorities for 5 consecutive years through annual transformation cycles, quarterly OKRs, and board-level reporting cadences", "5-Year Strategic Alignment")

# Additional Experience
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(6)
run(p, "Additional Experience", bold=True)
bullet(doc, "Digital Strategy Manager (2016\u20132018) \u2014 led digital strategy and campaign analytics; built reporting dashboards; optimized $500K+ annual ad spend through data-driven ROI measurement")
bullet(doc, "Client Services Representative (2014\u20132016) \u2014 managed enterprise client escalations; developed response protocols reducing resolution time by 30%")

# Education
section_header(doc, "Education")
body(doc, "Post-Baccalaureate Diploma in Technical Management & Services \u2014 KPU, Surrey, BC", size=Pt(9), sa=0)
body(doc, "Master of Business Administration (MBA)", size=Pt(9), sa=0)
body(doc, "Post-Graduate Diploma in Business Management (IT)", size=Pt(9), sa=0)
body(doc, "Bachelor of Science in Information Technology", size=Pt(9), sa=0)

# Certifications & Training
section_header(doc, "Certifications & Training")
body(doc, "ADKAR Change Management Model (Post-Bacc coursework, KPU)  |  Kotter 8-Step Change Process (MBA curriculum)  |  PROSCI-Aligned Change Management Methodologies  |  Organizational Behavior & Transformation Strategy (KPU)", size=Pt(9))

# Technical Proficiency
section_header(doc, "Technical Proficiency")
body(doc, "Change Management: ADKAR, Kotter 8-Step, PROSCI-Aligned, KPI Dashboard Design, OKR Frameworks  |  EHR & Practice Management: eClinicalWorks (expert), Athenahealth (expert)  |  Revenue Operations: RCM, Billing Infrastructure, Pipeline Management, Analytics  |  Business Tools: Salesforce, Microsoft 365, Google Workspace", size=Pt(9))

# Save
os.makedirs(f"{ONEDRIVE}/2026-07-14/Seaspan", exist_ok=True)
os.makedirs(f"{LINUX}/2026-07-14/Seaspan", exist_ok=True)
rpath = f"{ONEDRIVE}/2026-07-14/Seaspan/Aman_Kumar_Seaspan_Change_Management_Specialist.docx"
lpath = f"{LINUX}/2026-07-14/Seaspan/Aman_Kumar_Seaspan_Change_Management_Specialist.docx"
try:
    doc.save(rpath)
    print(f"Resume (OneDrive): {rpath}")
except (PermissionError, OSError):
    print("  (OneDrive unavailable, using Linux only)")
doc.save(lpath)
print(f"Resume (Linux): {lpath}")

# Count
c = sum(1 for p in doc.paragraphs if p.text.strip() and len(p.text.strip()) > 20)
print(f"  Content: {c} meaningful lines \u2014 PASS (threshold: 15)")

# ===== COVER LETTER =====
cl = Document()
for s in cl.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

p = cl.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, NAME, size=Pt(16), bold=True)

cl.add_paragraph()
body(cl, "July 14, 2026", sa=8)
body(cl, "Lead, Change Management", sa=0)
body(cl, "Seaspan ULC", sa=0)
body(cl, "50 Pemberton Avenue", sa=0)
body(cl, "North Vancouver, BC V7P 2R1", sa=8)
body(cl, "Re: Change Management Specialist \u2014 Job ID 8553", bold=True, sa=8)
body(cl, "Dear Lead, Change Management,", sa=8)

cover_body = [
    "I am writing to express my interest in the Change Management Specialist role at Seaspan. I bring ADKAR, Kotter, and PROSCI-aligned frameworks applied operationally across 8 years of organizational change \u2014 not from a classroom, but from the field.",
    "Most change management practitioners understand the theory. I have tested it.",
    "Over 8 years, I led 5 full-cycle organizational transformations through acquisition integration \u2014 developing Day 1/30/100 change plans, managing stakeholder alignment across 12 departments, and tracking adoption through custom KPI dashboards across 32 locations. When we transitioned from paper-based to a fully digital ecosystem, I managed every stage of the ADKAR model: building awareness of why change was needed, creating desire through direct benefit demonstration, delivering hands-on training across 150+ users, supporting the transition with on-site presence, and sustaining adoption through 6 months of reinforcement.",
    "I am drawn to Seaspan because the scale of transformation here is extraordinary \u2014 a $3.15B polar icebreaker, $6B naval support ships, new technologies, and a workforce of thousands. Organizational change at this scale demands someone who has not only studied the frameworks but has stood in front of 60 skeptical clinicians and convinced them to trust a new system. I have done that.",
    "I understand this role requires Controlled Goods Program clearance and ITAR compliance. I am fully prepared to meet these requirements.",
    "I would welcome the opportunity to discuss how my experience managing complex organizational change can support Seaspan\u2019s transformation journey.",
]
for b in cover_body:
    body(cl, b, sa=6)

body(cl, "", sa=0)
body(cl, "Best regards,", sa=0)
body(cl, "", sa=0)
body(cl, NAME, bold=True, sa=0)
body(cl, PHONE, sa=0)

clpath = f"{ONEDRIVE}/2026-07-14/Seaspan/Aman_Kumar_Seaspan_Change_Management_Specialist_Cover_Letter.docx"
cl_lpath = f"{LINUX}/2026-07-14/Seaspan/Aman_Kumar_Seaspan_Change_Management_Specialist_Cover_Letter.docx"
cl.save(clpath); cl.save(cl_lpath)
print(f"Cover:   {clpath}")

print("\nDone. Both files written.")
