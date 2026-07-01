#!/usr/bin/env python3
"""Generate DOCX for Clio, Deloitte, EY-Parthenon from existing SHOOT packages."""
import os, re, json, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(ROOT, "JOBS-OS-2026", "local_config.json")
if os.path.exists(CONFIG_PATH):
    with open(CONFIG_PATH) as f:
        cfg = json.load(f)
else:
    cfg = {}
REAL_NAME = cfg.get("name", "Aman Kumar")
REAL_PHONE = cfg.get("phone", "+1 236-885-2285")  
REAL_EMAIL = cfg.get("email", "amankumar7111@outlook.com")
REAL_LINKEDIN = cfg.get("linkedin", "https://www.linkedin.com/in/aman1776")

OUT_DIR = os.path.join(ROOT, "2026-06-30", "DOCX_OUTPUT")
os.makedirs(OUT_DIR, exist_ok=True)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for tag, val in [('w:color', '0000FF'), ('w:u', 'single')]:
        e = OxmlElement(tag)
        e.set(qn('w:val'), val)
        rPr.append(e)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20')
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def make_contact_parts(line):
    """Return list of (text, url_or_None) tuples from a contact line."""
    line = line.replace('[NAME]', REAL_NAME).replace('[PHONE]', REAL_PHONE)
    line = line.replace('[EMAIL]', REAL_EMAIL).replace('[LINKEDIN]', REAL_LINKEDIN)
    email_m = re.search(r'([\w.+-]+@[\w-]+\.[\w.]+)', line)
    li_m = re.search(r'(https?://(?:www\.)?linkedin\.com/in/[\w-]+)', line)
    simple_li = re.search(r'(www\.linkedin\.com/in/[\w-]+)', line)
    
    parts = []
    if email_m and li_m:
        add_text(line[:email_m.start()], parts)
        parts.append(('link', email_m.group(1), f'mailto:{email_m.group(1)}'))
        add_text(line[email_m.end():li_m.start()], parts)
        parts.append(('link', li_m.group(1), li_m.group(1)))
        add_text(line[li_m.end():], parts)
    elif email_m:
        add_text(line[:email_m.start()], parts)
        parts.append(('link', email_m.group(1), f'mailto:{email_m.group(1)}'))
        add_text(line[email_m.end():], parts)
    elif li_m:
        add_text(line[:li_m.start()], parts)
        parts.append(('link', li_m.group(1), li_m.group(1)))
        add_text(line[li_m.end():], parts)
    elif simple_li:
        url = f'https://{simple_li.group(1)}'
        add_text(line[:simple_li.start()], parts)
        parts.append(('link', simple_li.group(1), url))
        add_text(line[simple_li.end():], parts)
    else:
        add_text(line, parts)
    return parts

def add_text(t, parts):
    t = t.strip()
    if t:
        parts.append(('text', t, None))

def render_line(line, doc, is_cover=False):
    s = line.strip()
    if not s:
        doc.add_paragraph('').paragraph_format.space_after = Pt(2)
        return
    s = s.replace('[NAME]', REAL_NAME).replace('[PHONE]', REAL_PHONE)
    s = s.replace('[EMAIL]', REAL_EMAIL).replace('[LINKEDIN]', REAL_LINKEDIN)
    
    if is_cover and any(x in s for x in ['[NAME]', REAL_NAME, '@', 'linkedin.com', 'www.linkedin', '+1']):
        doc.add_paragraph(s).paragraph_format.space_after = Pt(2)
        return
    
    p = doc.add_paragraph()
    r = p.add_run(s)
    r.font.size = Pt(10) if not is_cover else Pt(10.5)
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    return p

def gen_resume(lines, save_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(10)
    for sec in doc.sections:
        sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
    
    for line in lines:
        s = line.strip()
        if not s:
            doc.add_paragraph('').paragraph_format.space_after = Pt(1)
            continue
        if s == REAL_NAME:
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.font.size = Pt(16); r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
        elif any(x in s.lower() for x in ['vancouver, bc', '@', 'linkedin', REAL_PHONE[:5], REAL_EMAIL[:5]]):
            p = doc.add_paragraph()
            for typ, text, url in make_contact_parts(s):
                if url:
                    add_hyperlink(p, text, url)
                else:
                    rr = p.add_run(text)
                    rr.font.size = Pt(9); rr.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
        elif s == s.upper() and len(s) > 3 and len(s) < 60:
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.font.size = Pt(11); r.bold = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            b = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
            bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '555555')
            b.append(bot); pPr.append(b)
        elif s.startswith('\u2022') or s.startswith('- '):
            txt = s.lstrip('\u2022- ').strip()
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            r = p.add_run(txt)
            r.font.size = Pt(10); r.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.3)
        elif any(w in s.lower() for w in ['director of', 'chief', 'manager,', 'lead ', 'head ']) and len(s) < 80:
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.font.size = Pt(10); r.bold = True
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
        else:
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.font.size = Pt(10); r.font.name = 'Calibri'
    doc.save(save_path)
    return save_path

def gen_cover(lines, save_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(10.5)
    for sec in doc.sections:
        sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)
    
    for line in lines:
        s = line.strip()
        if not s:
            doc.add_paragraph('').paragraph_format.space_after = Pt(4)
            continue
        s = s.replace('[NAME]', REAL_NAME).replace('[PHONE]', REAL_PHONE)
        s = s.replace('[EMAIL]', REAL_EMAIL).replace('[LINKEDIN]', REAL_LINKEDIN)
        if s.startswith('**') and ':' in s:
            continue
        p = doc.add_paragraph()
        r = p.add_run(s)
        r.font.size = Pt(10.5); r.font.name = 'Calibri'
        if any(x in s for x in ['Best regards', 'Sincerely', 'Warmly']):
            p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    doc.save(save_path)
    return save_path

# ── CLO ──
clio_resume = """Aman Kumar
amankumar7111@outlook.com | www.linkedin.com/in/aman1776 | +1 236-885-2285 | Vancouver, BC, Canada

PROFESSIONAL SUMMARY
Strategy and operations leader with 8 years of experience building and scaling multi-site operations from startup to $17M exit. Served as embedded strategic consultant to CEO — leading corporate strategy, OKR programs, business reviews, investment cases, and cross-functional execution across 32 locations. MBA-trained systems builder who turns ambiguity into operating rhythm. Seeking to bring BizOps expertise to a high-growth technology company transforming an industry.

CORE COMPETENCIES
Strategy & Operations | Corporate Strategy | OKR Programs | Business Reviews | Market Sizing | Financial Modeling | Competitive Analysis | Cross-Functional Leadership | Investment Case Development | Operating Rhythm Design | M&A Integration | P&L Management | AI Workflow Automation

PROFESSIONAL EXPERIENCE

SkyflyMD — Multi-site Healthcare Operations Group | Director of Operations
Arizona / Texas / Vancouver, BC | 2018–2026

• Led corporate strategy and annual planning — designed OKR framework adopted across 5 business units, facilitated quarterly business reviews with CEO and leadership team, built executive reporting cadence that tracked 50+ KPIs across 32 locations
• Built investment cases for growth — sized markets across 5 US states, modeled financial projections, assessed competitive landscape, and presented board-ready recommendations that guided expansion strategy
• Directed full-cycle M&A through $17M acquisition — structured 8 concurrent due diligence workstreams, built Day 1/100 integration playbook, consolidated 8 operational systems into unified platform, retained 100% key talent through earnout
• Designed company-wide operating rhythm — communication protocols, decision-making frameworks, meeting cadences, escalation paths — enabling 62% faster implementation of strategic initiatives
• Managed P&L of $4M ARR across 12 departments — forecasting, budgeting, vendor consolidation ($200K annual savings), board-level financial reporting
• Led technology transformation — implemented AI automation across 60% of operational workflows, achieving 99% verification accuracy and 40% improvement in service delivery efficiency
• Built organizational infrastructure from zero — defined role families, training programs, quality standards, and escalation paths for a remote cross-border team spanning India, US, and Canada
• Served as internal strategic consultant — led special projects on organizational design, process architecture, and growth strategy across 5 clinic groups

EDUCATION

MBA, Strategy & Operations
BSc, Information Technology
Post-Baccalaureate Diploma, Business Administration — KPU

TECHNICAL PROFICIENCY
Strategic Planning & OKR Frameworks | Financial Modeling & P&L Management | M&A Due Diligence & Integration | Business Intelligence & KPI Dashboards | AI Workflow Automation | Cross-Border Operations | Stakeholder Management & Executive Communication"""

clio_cover = """Aman Kumar
amankumar7111@outlook.com | www.linkedin.com/in/aman1776 | +1 236-885-2285
Vancouver, BC, Canada

Date: 30/06/2026

Clio Hiring Team
Clio (Themis Solutions Inc.)
Burnaby, BC

Re: Senior Strategy and Operations Associate (REQ-4583)

Dear Clio Hiring Team,

I've spent the last 8 years as an embedded strategy and operations leader — building the operational infrastructure, strategic planning function, and operating rhythm for a business that scaled from 3 to 70 people across 32 locations. When I read the description for Clio's Senior Strategy and Operations Associate role, I recognized the work immediately. It's the work I've been doing every day, and it's the work I want to do at Clio.

Clio's trajectory is remarkable. $500M in ARR. Clio Operate transforming the enterprise legal market. Clio Capital creating a new fintech category for law firms. The vLex and Jurisage acquisitions bringing AI-powered legal intelligence into the core platform. At this inflection point, the Strategy & Operations team — sitting at the center of the business, partnering with the C-Suite, Product, Corp Dev, and Finance — is not a support function. It's the connective tissue that turns strategy into execution.

I've played this exact role. At my previous company, I started with no playbook, no team, and no established processes. By the time we exited at $17M, I had built the corporate strategy function from scratch — OKR frameworks, quarterly business reviews, board-ready financial models, and an operating rhythm that turned ambiguity into accountability across 32 locations. I didn't inherit a playbook. I created one. This is what "Draw the Owl" means, and it's how I've operated my entire career.

The role calls for someone who can synthesize business insights, evaluate growth opportunities, and support strategy with analytical rigor. I've walked this bridge daily: translating operational data into strategic decisions, building investment cases for expansion across 5 US states, leading market sizing and competitive analysis, and facilitating the cross-functional collaboration that makes strategy real. My MBA training and hands-on P&L management give me the financial modeling and analytical skills. My years as the CEO's strategic partner give me the executive communication and stakeholder management that make analyses actionable.

AI is central to how I build. I led the automation of 60% of operational workflows, achieving 99% accuracy and 40% efficiency gains — not as a technical project, but as a strategic transformation that freed the team to focus on higher-value work. Clio's commitment to training 25,000 legal professionals on AI resonates deeply. I don't just follow AI trends — I've deployed AI to drive measurable business outcomes.

Clio's values are not just words on a page to me. "Win and Help Win" is how I built a team where 94% of employees stayed through an ownership change. "Get Better Every Day" is a continuous improvement system I designed and embedded into every process. "Lead with Context" is the governance model I built — transparent reporting, open communication, clear decision rights. I've lived these values at a similar stage of company growth.

I would welcome the opportunity to discuss how my experience building a strategy and operations function from scratch can support Clio's next chapter of growth.

Best regards,
Aman Kumar
Phone: +1 236-885-2285 | Email: amankumar7111@outlook.com | LinkedIn: www.linkedin.com/in/aman1776"""

# ── DELOITTE ──
deloitte_resume = """Aman Kumar
amankumar7111@outlook.com | www.linkedin.com/in/aman1776 | +1 236-885-2285 | Vancouver, BC, Canada

PROFESSIONAL SUMMARY
Operations executive who built a 70-person, 32-location organization from scratch, led digital transformation, managed P&L, and delivered $17M exit. Delivered strategic transformation programs across multi-site operations — technology enablement, operating model design, and governance frameworks. Brings first-hand program delivery experience to Deloitte's Consumer industry practice, where clients need leaders who have done the work, not just studied it.

CORE COMPETENCIES
Strategic Transformation | Technology Enablement | Program Delivery | P&L Management | Cross-Functional Leadership | Governance Frameworks | Consumer Industry Operations | Operational Excellence | M&A Integration | Organizational Design

PROFESSIONAL EXPERIENCE

SkyflyMD — Multi-site Healthcare Operations Group | Director of Operations
Arizona / Texas / Vancouver, BC | 2018–2026

• Led end-to-end strategic transformation across 32 locations — designed operating model, implemented technology stack, built governance framework, and established performance management systems from zero
• Managed P&L of $4M ARR across 12 departments — forecasting, budgeting, vendor consolidation ($200K annual savings), board-level financial reporting to CEO and investors
• Directed technology enablement program — implemented EHR, RCM, billing, analytics platforms across all locations, consolidated 8 legacy systems into unified stack, achieved 62% faster implementations
• Led cross-functional team of 12 department heads across 5 clinic groups — built operating rhythm, communication protocols, decision-making frameworks that scaled from 3 to 70 employees
• Drove $17M acquisition from due diligence through post-merger integration — structured 8 workstreams, built Day 1/100 playbook, retained 100% key talent, stabilized operations within 90 days
• Built governance frameworks across 5 jurisdictions (US and Canada) — compliance protocols, risk management, quality standards, escalation paths
• Served as primary interface between CEO and all clinic group managers — translated strategic direction into accountable execution plans
• Managed 10K+ consumer touchpoints across 32 locations — service delivery standards, customer experience metrics, quality assurance programs

EDUCATION

MBA, Strategy & Operations
BSc, Information Technology
Post-Baccalaureate Diploma, Business Administration — KPU

TECHNICAL PROFICIENCY
Strategic Transformation | Program Delivery | Technology Enablement | P&L Management | Governance Frameworks | Cross-Functional Leadership | Operational Excellence | M&A Integration | Consumer Industry Operations"""

deloitte_cover = """Aman Kumar
amankumar7111@outlook.com | www.linkedin.com/in/aman1776 | +1 236-885-2285
Vancouver, BC, Canada

Date: 30/06/2026

Deloitte Canada Recruitment Team
Vancouver, BC

Re: Senior Manager, Strategic & Technology Change – Consumer Industry

Dear Deloitte Team,

I built the business your consumer clients want to become. Over 8 years, I scaled a multi-site organization from 3 to 70 people across 32 locations, grew revenue to $4M ARR, led the technology transformation, managed the P&L, and directed a $17M exit. I'm applying to Deloitte because I've lived what your clients are trying to do. I don't need to study their business — I need to help make it work better.

Deloitte's Consumer practice helps clients navigate the most disruptive era in retail and consumer goods — changing customer expectations, supply chain reinvention, and digital-first operating models. These are the exact challenges I've solved operationally. I've managed 10,000+ consumer touchpoints across 32 locations, implemented enterprise technology systems from scratch, and built governance frameworks that ensured compliance across multiple jurisdictions.

The role calls for someone who can deliver strategic and technology change programs for consumer industry clients. My career is a case study in exactly this work. I didn't just advise on technology enablement — I led the implementation of EHR, RCM, billing, and analytics platforms across a distributed multi-site operation. I didn't just design operating models — I built one from zero that scaled from 3 to 70 people. I didn't just study P&L management — I managed $4M in revenue and $3M+ in expenses across 12 departments.

What sets me apart from a career consultant is depth. I've been inside a transformation, not adjacent to it. I know what works in practice, not just in theory. I know where programs go wrong — governance gaps, stakeholder misalignment, change management underinvestment — because I've recovered from those mistakes. For Deloitte's clients, that first-hand experience means faster delivery, fewer surprises, and better outcomes.

I would welcome the opportunity to discuss how my first-hand experience with strategic and technology transformation can contribute to Deloitte's Consumer practice in Vancouver.

Best regards,
Aman Kumar
Phone: +1 236-885-2285 | Email: amankumar7111@outlook.com | LinkedIn: www.linkedin.com/in/aman1776"""

# ── EY-PARTHENON ──
ey_resume = """Aman Kumar
amankumar7111@outlook.com | www.linkedin.com/in/aman1776 | +1 236-885-2285 | Vancouver, BC, Canada

PROFESSIONAL SUMMARY
Operations and strategy leader with 8 years of experience scaling multi-site organizations from startup to $17M exit. Directed the operational workstreams of a full-cycle acquisition — structuring due diligence across 8 functional areas, leading post-merger integration, and building the playbooks, teams, and systems that made the deal successful. Brings first-hand M&A execution experience to EY-Parthenon's Deal Management practice, where strategy meets delivery.

CORE COMPETENCIES
M&A Integration | Operational Due Diligence | Workstream Leadership | Post-Merger Integration | Synergy Realization | Integration Playbook Development | Cross-Functional Program Management | Executive Stakeholder Management | Organizational Design | Multi-Site Operations

PROFESSIONAL EXPERIENCE

SkyflyMD — Multi-site Healthcare Operations Group | Director of Operations
Arizona / Texas / Vancouver, BC | 2018–2026

• Directed operational workstreams for $17M acquisition — structured 8 concurrent due diligence tracks (finance, legal, operations, provider contracts, IT, HR, quality, facilities) and built integration playbook adopted company-wide
• Led post-merger integration: consolidated 8 operational systems to 1 unified stack, retained 100% of key talent through earnout period, achieved operational stability within 90 days of close
• Designed integration governance model with RAID log, escalation paths, and executive reporting cadence — enabled real-time visibility into workstream progress and risk
• Built M&A playbooks for functional departments — standardized approach to due diligence, Day 1 readiness, and post-close integration across all future acquisitions
• Scaled operations from 3 to 70 FTEs across 5 clinic groups — designed role families, training programs, quality standards, and multi-site operating rhythm from scratch
• Drove $200K annual savings via vendor consolidation program — built board-ready financial models and led cross-functional sourcing initiative
• Served as primary interface between CEO and all clinic group managers — translated strategic direction into accountable execution plans across 12 departments

EDUCATION

Master of Business Administration (MBA) — KPU
Post-Baccalaureate Diploma in Business Management — KPU
Bachelor of Science in Information Technology — First Class Honours

TECHNICAL PROFICIENCY
M&A Integration | Operational Due Diligence | Workstream Leadership | Post-Merger Integration | Synergy Realization | Executive Stakeholder Management | Organizational Design | Multi-Site Operations | Financial Modeling"""

ey_cover = """Aman Kumar
amankumar7111@outlook.com | www.linkedin.com/in/aman1776 | +1 236-885-2285
Vancouver, BC, Canada

Date: 30/06/2026

EY-Parthenon Recruitment Team
Vancouver, BC

Re: Senior Manager, Strategy & Execution – Deal Management (Requisition 1685558)

Dear EY-Parthenon Team,

I am writing to apply for the Senior Manager, Strategy & Execution – Deal Management position in Vancouver. I have spent the last eight years inside a multi-site healthcare services organization — scaling it from 3 employees to 70, from 1 location to 32, from zero systems to a consolidated operational infrastructure — and I directed the operational workstreams of our $17M acquisition from due diligence through post-merger integration. The end-to-end deal lifecycle I led maps directly to the work your Deal Management practice delivers every day.

I know how deals work on the ground — not just in theory. When our company was acquired, I designed and coordinated 8 concurrent due diligence workstreams covering finance, legal, operations, provider contracts, IT, HR, quality, and facilities. I built the RAID log, established the governance cadence, managed cross-functional dependencies, and surfaced risks to the CEO and buyer before they became deal issues. I developed the integration playbook that guided Day 1 readiness, systems consolidation, and the 90-day stabilization period post-close. Every one of those workstreams — operational due diligence, synergy tracking, integration planning, transition service agreements — I have done them.

I build systems that scale beyond the deal. Post-acquisition, I consolidated 8 separate operational systems into a single unified stack — standardizing processes across 5 clinic groups and 32 locations. I created playbooks for every department, designed role families and training programs, and established the operating cadence that enabled the combined organization to function as one. This is the integration playbook development and operating model design your role requires. I do not just advise on operating models — I build them.

I bring an execution mindset that consulting firms value. My background is not traditional consulting pedigree; it is building from scratch, directing through complexity, and getting results without a playbook. I managed 12 department budgets, reported directly to the CEO, led cross-functional initiatives across 5 distinct business units, and retained 95% of my team through one of the most disruptive events an organization can experience — an ownership change. That combination of operational depth and strategic thinking is rare in the deal management space.

I am drawn to EY-Parthenon specifically because the Strategy & Execution practice lives at the intersection of strategy and delivery — advising on the deal AND making it happen. That is exactly where I operate best. I would welcome the opportunity to discuss how my deal experience and operational expertise can contribute to your client engagements.

Thank you for your consideration. I look forward to hearing from you.

Best regards,
Aman Kumar
Phone: +1 236-885-2285 | Email: amankumar7111@outlook.com | LinkedIn: www.linkedin.com/in/aman1776"""

# ── GENERATE FILES ──
print("=" * 60)
print("  GENERATING DOCX FILES")
print("=" * 60)

packages = [
    ("Clio_Senior_Strategy_Operations_Associate", clio_resume, clio_cover),
    ("Deloitte_SrMgr_Strategic_Technology_Change", deloitte_resume, deloitte_cover),
    ("EY_Parthenon_SrMgr_Deal_Management", ey_resume, ey_cover),
]

for slug, rtext, ctext in packages:
    rp = os.path.join(OUT_DIR, f"Aman_{slug}_Resume_2026-06-30.docx")
    cp = os.path.join(OUT_DIR, f"Aman_{slug}_CoverLetter_2026-06-30.docx")
    gen_resume(rtext.split('\n'), rp)
    gen_cover(ctext.split('\n'), cp)
    print(f"  ✓ {slug}")
    print(f"    Resume: {rp}")
    print(f"    Cover:  {cp}")

print(f"\n  Output directory: {OUT_DIR}")
print("=" * 60)
