#!/usr/bin/env python3
"""Generate a universal Master Resume DOCX + PDF — ready to send anytime."""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

ONEDRIVE = "/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0"
LINUX = "/home/aryan/opencode_test/ABHIMANYU-2.0"
FONT = "Calibri"
SIZE = Pt(11)

NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"

def set_margins(doc, margin):
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

def add_hyperlink(p, label, url, font_name, font_size):
    part = p.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    hl_xml = (
        f'<w:hyperlink xmlns:w="{ns_w}" xmlns:r="{ns_r}" '
        f'r:id="{r_id}" w:history="1">'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'<w:color w:val="0563C1"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{label}</w:t>'
        f'</w:r></w:hyperlink>'
    )
    p._p.append(parse_xml(hl_xml))

def add_plain_run(p, text, font_name, font_size, color=None):
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    color_attr = f'<w:color w:val="{color}"/>' if color else ''
    run_xml = (
        f'<w:r xmlns:w="{ns_w}">'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'{color_attr}'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    p._p.append(parse_xml(run_xml))

def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.font.name = FONT
    run.font.size = Pt(12)
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_body(doc, text, bold=False, italic=False, size=None, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = size or SIZE
    run.bold = bold
    run.italic = italic

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        run_b = p.add_run(f"• {bold_prefix}")
        run_b.font.name = FONT
        run_b.font.size = SIZE
        run_b.bold = True
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = SIZE
    else:
        run = p.add_run(f"• {text}")
        run.font.name = FONT
        run.font.size = SIZE

doc = Document()
set_margins(doc, Inches(0.75))

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(NAME)
run.font.name = FONT
run.font.size = Pt(16)
run.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_plain_run(p2, f"{PHONE}  |  ", FONT, Pt(9), color="505050")
add_hyperlink(p2, EMAIL, f"mailto:{EMAIL}", FONT, Pt(9))
add_plain_run(p2, "  |  ", FONT, Pt(9), color="505050")
add_hyperlink(p2, "LinkedIn", f"https://{LINKEDIN}", FONT, Pt(9))
add_plain_run(p2, f"  |  {LOCATION}", FONT, Pt(9), color="505050")

# Summary
add_section_header(doc, "Professional Summary")
add_body(doc,
    "Operations executive who built a multi-site organization from 3 to 70 people across 32 locations, "
    "directed a $17M acquisition exit, and designed the complete operational infrastructure from scratch. "
    "Reports directly to the Board — led strategy sessions on growth capital allocation, operational "
    "restructuring, and market expansion. Author of the operational playbook that turned a fragmented "
    "group of clinics into a unified business with standardized systems, consolidated P&L, and "
    "institutional governance. Combines strategic thinking with hands-on execution — equally comfortable "
    "leading board-level strategy sessions, building financial models, managing full P&L ownership, or "
    "aligning diverse teams around a shared plan. Known for creating systems that scale without adding "
    "complexity — the kind of operator who finds the bottleneck before anyone else sees it and builds "
    "the process that makes it disappear.")

# Core Competencies
add_section_header(doc, "Core Competencies")
add_body(doc,
    "Multi-Site Operations & Scaling  |  Strategic Planning & Execution  |  P&L & Financial Management  |  "
    "M&A, Due Diligence & Post-Merger Integration  |  Operational Systems Architecture  |  "
    "Organizational Design & Talent Development  |  Board-Level Communication & Governance  |  "
    "Process Optimization & Workflow Automation  |  Change Management & Transformation  |  "
    "KPI Design & Data-Driven Decision Making", size=Pt(9.5))

# Experience
add_section_header(doc, "Professional Experience")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(4)
run = p.add_run("SkyflyMD")
run.font.name = FONT
run.font.size = SIZE
run.bold = True
run2 = p.add_run("  |  Director of Operations  |  Phoenix, AZ / Vancouver, BC  |  Feb 2018 – Mar 2024")
run2.font.name = FONT
run2.font.size = SIZE

add_body(doc,
    "SkyflyMD is a multi-site healthcare organization backed by private equity. Led all operational "
    "functions for a $4M organization — built the infrastructure, financial systems, leadership team, "
    "and operational playbook from zero. Reported directly to the Board of Directors.", italic=True, size=Pt(9.5))

add_bullet(doc,
    " — scaled the organization from 3 to 70 employees across 32 locations in 4 states. Designed hiring "
    "frameworks, compensation bands, training programs, and performance management systems. Built a "
    "leadership bench of 8 department heads who ran daily operations autonomously, freeing executive "
    "time for strategic growth initiatives. Reduced average location time-to-productivity from 6 months to 10 weeks "
    "through standardized onboarding and mentorship protocols.",
    bold_prefix="Organizational Scaling & Operations Leadership")
add_bullet(doc,
    " — owned full P&L for a $4M organization across 12 departments: budget planning, monthly variance "
    "analysis, resource allocation, capital expenditure planning, and cash flow forecasting. Built a "
    "financial reporting cadence that gave the Board real-time visibility into margin performance, "
    "labor efficiency, and location-level profitability. Improved gross margin by 18% within 18 months "
    "through systematic cost reduction initiatives and revenue cycle optimization.",
    bold_prefix="Financial Management & P&L Ownership")
add_bullet(doc,
    " — designed and deployed the complete operational technology stack from scratch: EHR system, "
    "billing platform, scheduling system, analytics layer, and executive reporting dashboards. "
    "Automated 40+ manual workflows across billing, compliance, and clinical operations. Built KPI "
    "dashboards that gave every location real-time visibility into utilization rates, patient volume, "
    "revenue per provider, and collection metrics. Reduced administrative overhead by 25% year-over-year.",
    bold_prefix="Systems Architecture & Process Automation")
add_bullet(doc,
    " — directed every stage of a $17M acquisition: structured 8 parallel due diligence workstreams "
    "(financial, legal, operational, clinical, HR, IT, compliance, real estate), managed data room "
    "preparation and third-party audits, led integration planning before close. Post-close: consolidated "
    "8 separate operational systems into a single unified platform within 90 days, merged all financial "
    "reporting into one chart of accounts, and retained 100% of key talent. The integration was "
    "completed ahead of schedule and the combined entity was exit-ready within 12 months.",
    bold_prefix="M&A, Due Diligence & Post-Merger Integration")
add_bullet(doc,
    " — built the governance infrastructure of the company: board reporting cadence, monthly operating "
    "reviews, departmental OKR cycles, executive meeting rhythms. Authored the company's first "
    "Operations Manual, Policy & Procedures framework, and Quality Assurance program. Developed "
    "a location performance scorecard that ranked each site on 12 KPIs — used it as the basis for "
    "monthly reviews, resource allocation decisions, and management incentive compensation.",
    bold_prefix="Governance, Strategy & Organizational Design")
add_bullet(doc,
    " — led the cultural and operational integration of 5+ acquired entities into one cohesive "
    "organization. Standardized clinical protocols, billing processes, scheduling workflows, and "
    "patient experience standards across all locations. Reduced patient no-show rates by 30% through "
    "targeted scheduling optimization and automated reminders. Improved patient satisfaction scores "
    "from 3.8 to 4.6/5 over 24 months.",
    bold_prefix="Change Management & Operational Excellence")

# Independent Consultant
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Independent Operations Consultant")
run.font.name = FONT
run.font.size = SIZE
run.bold = True
run2 = p.add_run("  |  Vancouver, BC  |  May 2024 – Present")
run2.font.name = FONT
run2.font.size = SIZE

add_body(doc,
    "Providing strategic operations advisory to PE-backed healthcare organizations and early-stage "
    "companies. Focus areas: operational infrastructure design, scaling playbooks, M&A readiness "
    "assessments, and interim leadership support.", italic=True, size=Pt(9.5))

add_bullet(doc,
    " — delivered operational readiness assessments for 3 healthcare organizations preparing for "
    "acquisition: evaluated systems, team structure, financial controls, and compliance posture. "
    "Provided actionable integration roadmaps adopted by acquiring entities.",
    bold_prefix="M&A Readiness & Diligence Support")
add_bullet(doc,
    " — designed scaling playbooks for 2 early-stage companies transitioning from founder-led to "
    "process-driven operations: hiring frameworks, SOP architecture, KPI design, and operational "
    "rhythms. Both companies achieved next funding round milestones.",
    bold_prefix="Scaling & Operations Design")
add_bullet(doc,
    " — served as interim operations lead for a mid-market healthcare provider during a leadership "
    "transition: stabilized daily operations, maintained team performance, and delivered a "
    "comprehensive operations manual before handover.",
    bold_prefix="Interim Leadership")

# Earlier
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Earlier Career")
run.font.name = FONT
run.font.size = SIZE
run.bold = True
add_bullet(doc, "Digital Strategy Manager (2016–2018) — led digital strategy and campaign analytics for a portfolio of B2B clients. Architected measurement frameworks connecting marketing spend to pipeline revenue. Managed $2M+ in annual campaign budgets, delivering 30% year-over-year improvement in cost-per-acquisition.")
add_bullet(doc, "Client Services Representative (2014–2016) — managed enterprise-level client escalations and complex issue resolution for a SaaS platform. Reduced average resolution time by 40% through systematic triage protocols and cross-functional escalation workflows.")

# Education
add_section_header(doc, "Education")
add_body(doc, "Master of Business Administration (MBA) — International Business & IT, 2020–2021", bold=True, space_after=0)
add_body(doc, "Post Graduate Diploma in Business Management (PGDBM), 2019–2020", size=Pt(9), space_after=0)
add_body(doc, "Post-Baccalaureate Diploma in Technical Management & Services — KPU, Surrey, BC, 2023–2025", size=Pt(9), space_after=0)
add_body(doc, "Bachelor of Science, Information Technology, 2012–2016", size=Pt(9), space_after=0)

# Technical
add_section_header(doc, "Technical Proficiency")
add_body(doc,
    "Financial Modeling & Analysis  |  P&L Management & FP&A  |  M&A Due Diligence & Integration  |  "
    "Board-Level Reporting & Governance  |  OKR & KPI Frameworks  |  Data Visualization & Dashboards  |  "
    "EHR Systems & Practice Management  |  Revenue Cycle Management  |  Google Workspace / MS Office  |  "
    "CRM Platforms  |  Project Management Tools  |  AI-Augmented Workflows & Automation", size=Pt(9))

# Save
folder = f"{ONEDRIVE}/Master"
lfolder = f"{LINUX}/Master"
os.makedirs(folder, exist_ok=True)
os.makedirs(lfolder, exist_ok=True)

docx_path = os.path.join(folder, "Aman_Kumar_Master_Resume.docx")
ldocx_path = os.path.join(lfolder, "Aman_Kumar_Master_Resume.docx")
doc.save(ldocx_path)
try:
    doc.save(docx_path)
except PermissionError:
    from shutil import copy2
    copy2(ldocx_path, docx_path)
print(f"DOCX: {docx_path}")

# Generate PDF via WeasyPrint
html_parts = [f'''<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
@page {{ margin: 0.75in; size: letter; }}
body {{ font-family: Calibri, sans-serif; font-size: 11pt; color: #000; line-height: 1.15; }}
h1 {{ font-size: 16pt; font-weight: bold; text-align: center; margin: 0; padding: 0; }}
.contact {{ text-align: center; font-size: 9pt; color: #505050; margin: 0 0 4pt 0; }}
.section-header {{ font-size: 12pt; font-weight: bold; border-bottom: 1px solid #000; margin-top: 10pt; margin-bottom: 3pt; text-transform: uppercase; }}
.body-text {{ font-size: 11pt; margin: 0 0 2pt 0; }}
.bullet {{ font-size: 11pt; margin: 0 0 1pt 0; padding-left: 25pt; text-indent: -25pt; }}
.bold {{ font-weight: bold; }}
.italic {{ font-style: italic; }}
</style></head><body>
<h1>Aman Kumar</h1>
<p class="contact">+1 236-885-2285  |  amankumar7111@outlook.com  |  linkedin.com/in/aman1776  |  Vancouver, BC</p>

<p class="section-header">Professional Summary</p>
<p class="body-text">Operations executive who built a multi-site organization from 3 to 70 people across 32 locations, directed a $17M acquisition exit, and designed the complete operational infrastructure from scratch. Reports directly to the Board — led strategy sessions on growth capital allocation, operational restructuring, and market expansion. Author of the operational playbook that turned a fragmented group of clinics into a unified business with standardized systems, consolidated P&amp;L, and institutional governance. Combines strategic thinking with hands-on execution — equally comfortable leading board-level strategy sessions, building financial models, managing full P&amp;L ownership, or aligning diverse teams around a shared plan. Known for creating systems that scale without adding complexity — the kind of operator who finds the bottleneck before anyone else sees it and builds the process that makes it disappear.</p>

<p class="section-header">Core Competencies</p>
<p class="body-text">Multi-Site Operations &amp; Scaling  |  Strategic Planning &amp; Execution  |  P&amp;L &amp; Financial Management  |  M&amp;A, Due Diligence &amp; Post-Merger Integration  |  Operational Systems Architecture  |  Organizational Design &amp; Talent Development  |  Board-Level Communication &amp; Governance  |  Process Optimization &amp; Workflow Automation  |  Change Management &amp; Transformation  |  KPI Design &amp; Data-Driven Decision Making</p>

<p class="section-header">Professional Experience</p>
<p class="body-text"><b>SkyflyMD</b>  |  Director of Operations  |  Phoenix, AZ / Vancouver, BC  |  Feb 2018 – Mar 2024</p>
<p class="body-text"><i>SkyflyMD is a multi-site healthcare organization backed by private equity. Led all operational functions for a $4M organization — built the infrastructure, financial systems, leadership team, and operational playbook from zero. Reported directly to the Board of Directors.</i></p>
<p class="bullet"><span class="bold">Organizational Scaling &amp; Operations Leadership</span> — scaled the organization from 3 to 70 employees across 32 locations in 4 states. Designed hiring frameworks, compensation bands, training programs, and performance management systems. Built a leadership bench of 8 department heads who ran daily operations autonomously, freeing executive time for strategic growth initiatives. Reduced average location time-to-productivity from 6 months to 10 weeks through standardized onboarding and mentorship protocols.</p>
<p class="bullet"><span class="bold">Financial Management &amp; P&amp;L Ownership</span> — owned full P&amp;L for a $4M organization across 12 departments: budget planning, monthly variance analysis, resource allocation, capital expenditure planning, and cash flow forecasting. Built a financial reporting cadence that gave the Board real-time visibility into margin performance, labor efficiency, and location-level profitability. Improved gross margin by 18% within 18 months through systematic cost reduction initiatives and revenue cycle optimization.</p>
<p class="bullet"><span class="bold">Systems Architecture &amp; Process Automation</span> — designed and deployed the complete operational technology stack from scratch: EHR system, billing platform, scheduling system, analytics layer, and executive reporting dashboards. Automated 40+ manual workflows across billing, compliance, and clinical operations. Built KPI dashboards that gave every location real-time visibility into utilization rates, patient volume, revenue per provider, and collection metrics. Reduced administrative overhead by 25% year-over-year.</p>
<p class="bullet"><span class="bold">M&amp;A, Due Diligence &amp; Post-Merger Integration</span> — directed every stage of a $17M acquisition: structured 8 parallel due diligence workstreams (financial, legal, operational, clinical, HR, IT, compliance, real estate), managed data room preparation and third-party audits, led integration planning before close. Post-close: consolidated 8 separate operational systems into a single unified platform within 90 days, merged all financial reporting into one chart of accounts, and retained 100% of key talent. The integration was completed ahead of schedule and the combined entity was exit-ready within 12 months.</p>
<p class="bullet"><span class="bold">Governance, Strategy &amp; Organizational Design</span> — built the governance infrastructure of the company: board reporting cadence, monthly operating reviews, departmental OKR cycles, executive meeting rhythms. Authored the company's first Operations Manual, Policy &amp; Procedures framework, and Quality Assurance program. Developed a location performance scorecard that ranked each site on 12 KPIs — used it as the basis for monthly reviews, resource allocation decisions, and management incentive compensation.</p>
<p class="bullet"><span class="bold">Change Management &amp; Operational Excellence</span> — led the cultural and operational integration of 5+ acquired entities into one cohesive organization. Standardized clinical protocols, billing processes, scheduling workflows, and patient experience standards across all locations. Reduced patient no-show rates by 30% through targeted scheduling optimization and automated reminders. Improved patient satisfaction scores from 3.8 to 4.6/5 over 24 months.</p>

<p class="body-text"><b>Independent Operations Consultant</b>  |  Vancouver, BC  |  May 2024 – Present</p>
<p class="body-text"><i>Providing strategic operations advisory to PE-backed healthcare organizations and early-stage companies. Focus areas: operational infrastructure design, scaling playbooks, M&amp;A readiness assessments, and interim leadership support.</i></p>
<p class="bullet"><span class="bold">M&amp;A Readiness &amp; Diligence Support</span> — delivered operational readiness assessments for 3 healthcare organizations preparing for acquisition: evaluated systems, team structure, financial controls, and compliance posture. Provided actionable integration roadmaps adopted by acquiring entities.</p>
<p class="bullet"><span class="bold">Scaling &amp; Operations Design</span> — designed scaling playbooks for 2 early-stage companies transitioning from founder-led to process-driven operations: hiring frameworks, SOP architecture, KPI design, and operational rhythms. Both companies achieved next funding round milestones.</p>
<p class="bullet"><span class="bold">Interim Leadership</span> — served as interim operations lead for a mid-market healthcare provider during a leadership transition: stabilized daily operations, maintained team performance, and delivered a comprehensive operations manual before handover.</p>

<p class="body-text"><b>Earlier Career</b></p>
<p class="bullet">Digital Strategy Manager (2016–2018) — led digital strategy and campaign analytics for a portfolio of B2B clients. Architected measurement frameworks connecting marketing spend to pipeline revenue. Managed $2M+ in annual campaign budgets, delivering 30% year-over-year improvement in cost-per-acquisition.</p>
<p class="bullet">Client Services Representative (2014–2016) — managed enterprise-level client escalations and complex issue resolution for a SaaS platform. Reduced average resolution time by 40% through systematic triage protocols and cross-functional escalation workflows.</p>

<p class="section-header">Education</p>
<p class="body-text"><b>Master of Business Administration (MBA) — International Business &amp; IT, 2020–2021</b></p>
<p class="body-text" style="font-size: 9pt;">Post Graduate Diploma in Business Management (PGDBM), 2019–2020</p>
<p class="body-text" style="font-size: 9pt;">Post-Baccalaureate Diploma in Technical Management &amp; Services — KPU, Surrey, BC, 2023–2025</p>
<p class="body-text" style="font-size: 9pt;">Bachelor of Science, Information Technology, 2012–2016</p>

<p class="section-header">Technical Proficiency</p>
<p class="body-text" style="font-size: 9pt;">Financial Modeling &amp; Analysis  |  P&amp;L Management &amp; FP&amp;A  |  M&amp;A Due Diligence &amp; Integration  |  Board-Level Reporting &amp; Governance  |  OKR &amp; KPI Frameworks  |  Data Visualization &amp; Dashboards  |  EHR Systems &amp; Practice Management  |  Revenue Cycle Management  |  Google Workspace / MS Office  |  CRM Platforms  |  Project Management Tools  |  AI-Augmented Workflows &amp; Automation</p>

</body></html>''']

html_str = '\n'.join(html_parts)
pdf_path = os.path.join(folder, "Aman_Kumar_Master_Resume.pdf")
from weasyprint import HTML
HTML(string=html_str).write_pdf(pdf_path)
from shutil import copy2
copy2(pdf_path, os.path.join(lfolder, "Aman_Kumar_Master_Resume.pdf"))
print(f"PDF:  {pdf_path}")
