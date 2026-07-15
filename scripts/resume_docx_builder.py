#!/usr/bin/env python3
"""Universal DOCX/PDF resume builder from .md source.

Reads a markdown resume file, parses by section headers, applies per-company
formatting config, generates ATS-optimized DOCX (always) + PDF (if configured).

Usage:
  python3 resume_docx_builder.py --source Master_Resume.md --company MASTER
  python3 resume_docx_builder.py --source Master_Resume.md --company Clio --role "Operations Lead"
  python3 resume_docx_builder.py --package 01_Company.md --company CompanyName
"""

import sys
import os
import re
import argparse
import subprocess
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

# --- CONSTANTS ---
ONEDRIVE = "/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0"
LINUX = "/home/aryan/opencode_test/ABHIMANYU-2.0"

NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"

# --- ATS PLATFORM -> FORMAT PREFERENCE ---
# Based on ATS_ESOTERICA.md lines 338-357
ATS_FORMAT = {
    "greenhouse": "pdf",
    "workday": "docx",
    "lever": "pdf",
    "ashby": "docx",
    "smartrecruiters": "docx",
    "bamboohr": "docx",
    "breezy": "pdf",
    "personio": "docx",
    "recruitee": "docx",
    "teamtailor": "pdf",
    "workable": "docx",
    "rippling": "docx",
    "avature": "docx",
    "pinpoint": "pdf",
    "jibeapply": "docx",
    "comeet": "pdf",
    "oracle_cloud": "docx",
    "ultipro": "docx",
    "icims": "docx",
    "sap_successfactors": "docx",
}

# --- PER-COMPANY CONFIG ---
CONFIG = {
    "MASTER": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": None,
        "format": "docx",
        "ats_notes": "Master resume for executive search firms — Calibri 11pt, 2-page"
    },
    "Clio": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": "lever",
        "format": "pdf",
        "ats_notes": "Clio Lever ATS — Calibri, PDF preferred per platform"
    },
    "1Password": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": "ashby",
        "format": "docx",
        "ats_notes": "1Password Ashby ATS — Calibri, DOCX"
    },
    "Arc'teryx": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": "workday",
        "format": "docx",
        "ats_notes": "Arc'teryx Workday — DOCX"
    },
    "Human Agency": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": "greenhouse",
        "format": "pdf",
        "ats_notes": "Human Agency Greenhouse — PDF preferred"
    },
    "EY-Parthenon": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": "workday",
        "format": "docx",
        "ats_notes": "EY-Parthenon Workday — Calibri 10pt, DOCX"
    },
    "Deloitte": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": "sap_successfactors",
        "format": "docx",
        "ats_notes": "Deloitte SAP SuccessFactors — DOCX"
    },
    "Methanex": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": None,
        "format": "docx",
        "ats_notes": "Methanex career portal — DOCX"
    },
    "Seaspan": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": "oracle_cloud",
        "format": "docx",
        "ats_notes": "Seaspan Oracle Cloud ATS — DOCX, Calibri 11pt, CM framing"
    },
    "KPMG": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_platform": "workday",
        "format": "docx",
        "ats_notes": "KPMG Workday — DOCX"
    },
}

DEFAULT_CONFIG = {
    "font": "Calibri",
    "size": Pt(11),
    "header_size": Pt(13),
    "margins": Inches(0.75),
    "pages": 2,
    "ats_platform": None,
    "format": "docx",
    "ats_notes": "Default ATS config — Calibri 11pt, DOCX"
}

SECTION_HEADERS = [
    "PROFESSIONAL SUMMARY",
    "CORE COMPETENCIES",
    "PROFESSIONAL EXPERIENCE",
    "EDUCATION",
    "TECHNICAL PROFICIENCY",
    "ADDITIONAL EXPERIENCE",
]


# ============================================================
# FORMATTING ENGINE (copied from gen_docx.py)
# ============================================================

def set_margins(doc, margin):
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def add_contact(doc, config, name=NAME, phone=PHONE, email=EMAIL, linkedin=LINKEDIN, location=LOCATION):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.name = config["font"]
    run.font.size = Pt(16)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_plain_run(p2, f"{phone}  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(p2, email, f"mailto:{email}", config["font"], Pt(9))
    add_plain_run(p2, "  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(p2, "LinkedIn", f"https://{linkedin}", config["font"], Pt(9))
    add_plain_run(p2, f"  |  {location}", config["font"], Pt(9), color="505050")


def add_section_header(doc, text, config):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.font.name = config["font"]
    run.font.size = config["header_size"]
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body(doc, text, config, bold=False, italic=False, size=None, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = config["font"]
    run.font.size = size or config["size"]
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text, config, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        run_b = p.add_run(f"\u2022 {bold_prefix}: ")
        run_b.font.name = config["font"]
        run_b.font.size = config["size"]
        run_b.bold = True
        run = p.add_run(text)
        run.font.name = config["font"]
        run.font.size = config["size"]
    else:
        run = p.add_run(f"\u2022 {text}")
        run.font.name = config["font"]
        run.font.size = config["size"]


def add_hyperlink_contact(p, label, url, font_name, font_size):
    part = p.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    hyperlink_xml = (
        f'<w:hyperlink xmlns:w="{ns_w}" xmlns:r="{ns_r}" '
        f'r:id="{r_id}" w:history="1">'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'<w:color w:val="0563C1"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{label}</w:t>'
        f'</w:r></w:hyperlink>'
    )
    hyperlink_elem = parse_xml(hyperlink_xml)
    p._p.append(hyperlink_elem)


def add_plain_run(p, text, font_name, font_size, color=None):
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    color_attr = f'<w:color w:val="{color}"/>' if color else ''
    run_xml = (
        f'<w:r xmlns:w="{ns_w}">'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'{color_attr}'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    run_elem = parse_xml(run_xml)
    p._p.append(run_elem)


def add_signature(doc, config):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    add_plain_run(p, "Best regards,", config["font"], config["size"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    add_plain_run(p2, NAME, config["font"], config["size"])
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)
    add_plain_run(p3, PHONE, config["font"], config["size"], color="505050")
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(0)
    p4.paragraph_format.space_before = Pt(0)
    add_hyperlink_contact(p4, EMAIL, f"mailto:{EMAIL}", config["font"], config["size"])


def count_docx_content(doc):
    """Count meaningful content paragraphs in docx."""
    count = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and len(text) > 20:
            count += 1
    return count


# ============================================================
# SOURCE PARSER
# ============================================================

def parse_resume_source(filepath):
    """Parse a resume .md file into structured sections.

    Returns dict with keys: name, contact, summary, competencies,
    experience (list of dicts), education (list), tech_proficiency,
    and raw sections dict.
    """
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()

    # Strip code fences if present
    text = re.sub(r'```[\w]*\n?', '', text)

    lines = text.split('\n')

    # Extract name (first non-empty line)
    name = ""
    contact = ""
    for line in lines:
        stripped = line.strip()
        if stripped and not name:
            name = stripped
        elif stripped and name and not contact:
            contact = stripped
        elif stripped:
            break

    # Split into sections by header
    sections = {}
    current_header = None
    current_lines = []

    header_pattern = re.compile(r'^(' + '|'.join(re.escape(h) for h in SECTION_HEADERS) + r')$')

    for line in lines[2:]:  # skip name and contact lines
        stripped = line.strip()
        if header_pattern.match(stripped.upper()) and (not current_lines or stripped in SECTION_HEADERS):
            if current_header:
                sections[current_header] = '\n'.join(current_lines).strip()
            current_header = stripped.upper()
            current_lines = []
        elif current_header:
            current_lines.append(line)

    if current_header:
        sections[current_header] = '\n'.join(current_lines).strip()

    # Parse Professional Experience into structured blocks
    experience = []
    exp_text = sections.get("PROFESSIONAL EXPERIENCE", "")
    if exp_text:
        experience = parse_experience_blocks(exp_text)

    # Parse Education into list
    education = []
    edu_text = sections.get("EDUCATION", "")
    if edu_text:
        education = [line.strip() for line in edu_text.split('\n') if line.strip()]

    # Parse Additional Experience if exists
    additional_experience = []
    addl_text = sections.get("ADDITIONAL EXPERIENCE", "")
    if addl_text:
        additional_experience = [line.strip() for line in addl_text.split('\n') if line.strip()]

    return {
        "name": name,
        "contact": contact,
        "summary": sections.get("PROFESSIONAL SUMMARY", ""),
        "competencies": sections.get("CORE COMPETENCIES", ""),
        "experience": experience,
        "education": education,
        "tech_proficiency": sections.get("TECHNICAL PROFICIENCY", ""),
        "additional_experience": additional_experience,
        "_raw_sections": sections,
    }


def parse_experience_blocks(text):
    """Parse Professional Experience section into job blocks.

    Each block: Company | Title | Location | Dates
                 description paragraph (optional)
                 - bullet
                 - bullet
    """
    blocks = []
    lines = text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Detect job header: Company | Title | Location | Dates (3 pipe separators)
        if stripped.count('|') == 3:
            parts = [p.strip() for p in stripped.split('|')]
            job = {
                "company": parts[0] if len(parts) > 0 else "",
                "title": parts[1] if len(parts) > 1 else "",
                "location": parts[2] if len(parts) > 2 else "",
                "dates": parts[3] if len(parts) > 3 else "",
                "description": "",
                "bullets": [],
            }

            i += 1
            # Skip blank lines after header
            while i < len(lines) and not lines[i].strip():
                i += 1

            # Read description (non-bullet paragraphs until first bullet)
            desc_lines = []
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('- '):
                desc_lines.append(lines[i].strip())
                i += 1

            if desc_lines:
                job["description"] = ' '.join(desc_lines)

            # Read bullets (skip blank lines first)
            while i < len(lines) and not lines[i].strip():
                i += 1
            bullets = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                bullet_text = lines[i].strip()[2:]  # remove "- "
                # Detect bold prefix: text before first " :" or ": "
                # Pattern: "Prefix: body" or "Prefix : body"
                bold_prefix = None
                body_text = bullet_text

                colon_match = re.match(r'^(.+?)\s*:\s*(.*)', bullet_text)
                if colon_match:
                    bold_prefix = colon_match.group(1).strip()
                    body_text = colon_match.group(2).strip()

                bullets.append({
                    "bold_prefix": bold_prefix,
                    "text": body_text,
                    "raw": bullet_text,
                })
                i += 1

            job["bullets"] = bullets
            blocks.append(job)
        else:
            i += 1

    return blocks


# ============================================================
# RESUME BUILDER
# ============================================================

def build_resume(data, config, company, role=None):
    """Build DOCX from parsed resume data."""
    doc = Document()
    set_margins(doc, config["margins"])

    # Name + Contact
    name = data.get("name", NAME)
    add_contact(doc, config, name=name)

    # Professional Summary
    if data.get("summary"):
        add_section_header(doc, "Professional Summary", config)
        add_body(doc, data["summary"], config)

    # Core Competencies
    if data.get("competencies"):
        add_section_header(doc, "Core Competencies", config)
        # Render at slightly smaller size
        add_body(doc, data["competencies"], config, size=Pt(9.5))

    # Professional Experience
    if data.get("experience"):
        add_section_header(doc, "Professional Experience", config)

        for job in data["experience"]:
            # Job header: Company | Title | Location | Dates
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(job["company"])
            run.font.name = config["font"]
            run.font.size = config["size"]
            run.bold = True

            title_str = f"  |  {job['title']}  |  {job['location']}  |  {job['dates']}"
            run2 = p.add_run(title_str)
            run2.font.name = config["font"]
            run2.font.size = config["size"]

            # Description paragraph (italic context)
            if job.get("description"):
                add_body(doc, job["description"], config, italic=True, size=Pt(9.5))

            # Bullets
            for bullet in job["bullets"]:
                add_bullet(doc, bullet["text"], config, bold_prefix=bullet.get("bold_prefix"))

    # Additional Experience
    addl = data.get("additional_experience", [])
    if addl:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("Additional Experience")
        run.font.name = config["font"]
        run.font.size = config["size"]
        run.bold = True
        for entry in addl:
            add_bullet(doc, entry, config)

    # Education
    if data.get("education"):
        add_section_header(doc, "Education", config)
        for edu_line in data["education"]:
            if edu_line:
                add_body(doc, edu_line, config, size=Pt(9), space_after=0)

    # Technical Proficiency
    if data.get("tech_proficiency"):
        add_section_header(doc, "Technical Proficiency", config)
        add_body(doc, data["tech_proficiency"], config, size=Pt(9))

    return doc


def build_cover_letter(data, config, company):
    """Build a generic cover letter DOCX for the master resume."""
    doc = Document()
    set_margins(doc, config["margins"])

    add_body(doc, NAME, config, bold=True, space_after=0)
    cover_contact = doc.add_paragraph()
    cover_contact.paragraph_format.space_after = Pt(0)
    cover_contact.paragraph_format.space_before = Pt(0)
    add_plain_run(cover_contact, f"{PHONE}  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(cover_contact, EMAIL, f"mailto:{EMAIL}", config["font"], Pt(9))
    add_plain_run(cover_contact, "  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(cover_contact, "LinkedIn", f"https://{LINKEDIN}", config["font"], Pt(9))

    today = datetime.now().strftime("%B %d, %Y")
    add_body(doc, today, config, space_after=8)

    add_body(doc, f"{company} Hiring Team", config, space_after=0)
    add_body(doc, LOCATION, config, space_after=8)

    role_str = "Operations Executive"
    add_body(doc, f"Re: {role_str}", config, bold=True, space_after=8)

    body = (
        f"Dear Hiring Team,\n\n"
        f"I am an operations executive who builds scalable infrastructure. Over 8 years, I built the "
        f"centralized operations backbone for a multi-site organization "
        f"across 32 locations — growing it from 3 to 70 employees, from $0 to $4M ARR, and directing "
        f"a $17M acquisition to exit. I own the full stack: strategy, P&L, M&A, technology "
        f"transformation, and organizational design.\n\n"
        f"I am currently seeking operations, strategy, or Chief of Staff roles at "
        f"organizations that need a builder who has already walked the path from startup to scale. "
        f"I welcome opportunities across technology, professional services, healthcare, and corporate "
        f"strategy — in Vancouver, BC, or remote Canada.\n\n"
        f"Please find my resume attached. I look forward to discussing how my experience building and "
        f"scaling from zero can deliver value for your organization.\n\n"
        f"Best regards,\n"
        f"Aman Kumar"
    )

    add_body(doc, body, config, space_after=0)
    add_body(doc, "", config, space_after=0)
    add_signature(doc, config)

    return doc


# ============================================================
# PDF CONVERSION
# ============================================================

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF using LibreOffice headless. Returns PDF path or None."""
    pdf_path = docx_path.replace('.docx', '.pdf')
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
             os.path.dirname(pdf_path), docx_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and os.path.exists(pdf_path):
            return pdf_path
        else:
            print(f"  PDF conversion warning: {result.stderr.strip()}")
            return None
    except FileNotFoundError:
        print("  PDF conversion skipped: LibreOffice not installed.")
        print("  Install: sudo apt install libreoffice")
        print("  Or manually: Word -> File -> Export -> Create PDF/XPS")
        return None
    except subprocess.TimeoutExpired:
        print("  PDF conversion timed out (>60s). Skipping.")
        return None
    except Exception as e:
        print(f"  PDF conversion error: {e}")
        return None


# ============================================================
# MAIN
# ============================================================

def build_docx_path(company, role=None, date_str=None):
    """Build OneDrive and Linux output paths."""
    if not date_str:
        date_str = datetime.now().strftime("%Y-%m-%d")
    role_str = role.replace(" ", "_") if role else "Master_Resume"
    folder_name = company.replace(" ", "_").replace("'", "")

    folder = f"{ONEDRIVE}/{date_str}/{folder_name}"
    lfolder = f"{LINUX}/{date_str}/{folder_name}"
    os.makedirs(folder, exist_ok=True)
    os.makedirs(lfolder, exist_ok=True)

    safe_company = company.replace(" ", "_").replace("'", "")
    respath = f"{folder}/Aman_Kumar_{safe_company}_{role_str}.docx"
    lpath = f"{lfolder}/Aman_Kumar_{safe_company}_{role_str}.docx"
    clpath = f"{folder}/Cover_Letter_{safe_company}_{role_str}.docx"
    cl_lpath = f"{lfolder}/Cover_Letter_{safe_company}_{role_str}.docx"

    return respath, lpath, clpath, cl_lpath


def main():
    parser = argparse.ArgumentParser(description="Generate ATS-optimized resume DOCX from .md source")
    parser.add_argument("--source", help="Path to resume .md file")
    parser.add_argument("--package", help="Path to SHOOT package .md file")
    parser.add_argument("--company", default="MASTER", help="Company name (for config, path, naming)")
    parser.add_argument("--role", default="Master_Resume", help="Role/title for file naming")
    parser.add_argument("--pdf", action="store_true", help="Also generate PDF (overrides config default)")
    args = parser.parse_args()

    if not args.source and not args.package:
        print("Error: Must provide --source or --package")
        sys.exit(1)

    # Resolve config
    config = CONFIG.get(args.company, DEFAULT_CONFIG)

    # Determine if PDF is needed
    generate_pdf = args.pdf or (config.get("format") == "pdf")

    # Parse source
    source_path = args.source or args.package
    print(f"Reading: {source_path}")

    data = parse_resume_source(source_path)
    if not data.get("summary") and not data.get("experience"):
        print(f"Error: Could not parse resume sections from {source_path}")
        print("  Check that section headers match: PROFESSIONAL SUMMARY, CORE COMPETENCIES, etc.")
        sys.exit(1)

    # Build resume DOCX
    doc = build_resume(data, config, args.company, args.role)
    respath, lpath, clpath, cl_lpath = build_docx_path(args.company, args.role)

    # Save resume (try OneDrive first, fallback to Linux)
    try:
        doc.save(respath)
    except (PermissionError, OSError):
        respath = lpath
        doc.save(respath)
    doc.save(lpath)
    print(f"Resume: {respath}")

    # Content integrity check
    c = count_docx_content(doc)
    if c < 15:
        print(f"FATAL: Resume only has {c} meaningful lines (< 15 minimum). Content too thin.")
        sys.exit(1)
    print(f"  Content: {c} meaningful lines (threshold: 15) — PASS")

    # Build cover letter DOCX
    cl_doc = build_cover_letter(data, config, args.company)
    try:
        cl_doc.save(clpath)
    except (PermissionError, OSError):
        clpath = cl_lpath
        cl_doc.save(clpath)
    cl_doc.save(cl_lpath)
    print(f"Cover:   {clpath}")

    # PDF conversion if needed
    if generate_pdf:
        print("  PDF mode enabled — converting...")
        # Resume PDF
        pdf_path = convert_to_pdf(respath)
        if pdf_path:
            print(f"  PDF:     {pdf_path}")
        # Cover letter PDF
        cl_pdf_path = convert_to_pdf(clpath)
        if cl_pdf_path:
            print(f"  Cover PDF: {cl_pdf_path}")

    print("\nDone.")


if __name__ == "__main__":
    main()
